import os
import re
import docx
import PyPDF2
from PyPDF2 import PdfReader
from docx import Document

def extract_name(text):
    words = text.strip().split()
    name = " ".join(words[:3]) if len(words) >= 3 else " ".join(words)
    return name


def extract_phone_number(text):
    # Regex pattern to match Indian phone numbers starting with +91
    phone_pattern = r'(\+91[\s-]?\d{10})'  # Matches +91 followed by 10 digits
    match = re.search(phone_pattern, text)
    return match.group(0) if match else "Not Found"

def extract_text_from_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

def extract_email(text):
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else "Not Found"

def extract_links_from_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        links = []

        for page in reader.pages:
            if '/Annots' in page:
                annotations = page['/Annots']
                for annot in annotations:
                    obj = annot.get_object()
                    if '/A' in obj and '/URI' in obj['/A']:
                        uri = obj['/A']['/URI']
                        if 'mailto:' in uri:
                            links.append(uri.replace('mailto:', ''))
        return links
    except Exception as e:
        print(f"Error extracting links: {e}")
        return []

def extract_full_name_combined(resume_text):
    lines = [line.strip() for line in resume_text.split("\n") if line.strip()]
    top_lines = lines[:5]

    name_parts = []
    for line in top_lines:
        if re.search(r'\d|@|http|linkedin|github|gmail', line, re.IGNORECASE): break
        if re.match(r'^[A-Za-z ]+$', line):  # Allow multi-word names
            name_parts.append(line)
        else:
            break

    return " ".join(name_parts).strip() if name_parts else "Unknown"

def parse_resume(resume_path):
    file_extension = os.path.splitext(resume_path)[1].lower()
    raw_text = ""

    if file_extension == '.docx':
        # Process DOCX file
        doc = docx.Document(resume_path)
        raw_text = " ".join([para.text for para in doc.paragraphs])
    elif file_extension == '.pdf':
        # Process PDF file
        with open(resume_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                raw_text += page.extract_text()
    else:
        raise ValueError("Unsupported file type. Only .docx and .pdf are supported.")

    # Extract name, email, and phone number from raw_text
    name = extract_full_name_combined(raw_text)
    email = extract_email(raw_text)
    emails_from_links = extract_links_from_pdf(resume_path)  # Updated here
    phone_number = extract_phone_number(raw_text)
    all_emails = emails_from_links + [email]  # Changed to make email a list
    email = all_emails[0] if all_emails else "Not Found"  # Extract phone number with +91

    # Return parsed data
    return {"raw_text": raw_text, "name": name, "email": email, "phone_number": phone_number}
